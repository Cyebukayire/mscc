import os
import requests
# from google.oauth2.credentials import Credentials
# from google_auth_oauthlib.flow import InstalledAppFlow
# from google.auth.transport.requests import Request
# from googleapiclient.discovery import build
# from googleapiclient.http import MediaFileUpload
from docx import Document
import pandas as pd
# from config.config import settings

# Paths
# CREDENTIALS_FILE = f'{settings.CONFIG_PATH}credentials.json'
# TOKEN_FILE = 'token.json'  # Path to save the access token
# SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/documents.readonly']


# Download word document from url
def download_docx_from_url(url, file_path):
    response = requests.get(url)
    with open(file_path, 'wb') as f:
        f.write(response.content)
        
def download_file(url, local_filename):
    with requests.get(url, stream=True) as r:
        r.raise_for_status()
        with open(local_filename, 'wb') as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)
    return local_filename

# def authenticate():
#     creds = None
#     if os.path.exists(TOKEN_FILE):
#         creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
#     if not creds or not creds.valid:
#         if creds and creds.expired and creds.refresh_token:
#             creds.refresh(Request())
#         else:
#             flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
#             creds = flow.run_local_server(port=0)
#         with open(TOKEN_FILE, 'w') as token:
#             token.write(creds.to_json())
#     return creds

# creds = authenticate()
# drive_service = build('drive', 'v3', credentials=creds)
# docs_service = build('docs', 'v1', credentials=creds)

# def upload_to_drive(filepath):
#     file_metadata = {'name': os.path.basename(filepath), 'mimeType': 'application/vnd.google-apps.document'}
#     media = MediaFileUpload(filepath, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
#     return file.get('id')

# def count_words_in_doc(doc_id):
#     doc = docs_service.documents().get(documentId=doc_id).execute()
#     text = ''
#     for element in doc.get('body').get('content'):
#         if 'paragraph' in element:
#             for paragraph_element in element.get('paragraph').get('elements'):
#                 if 'textRun' in paragraph_element:
#                     text += paragraph_element.get('textRun').get('content')
#     return len(text.split())

def count_words_with_docx_lib(temp_docx_file_path):
    doc = Document(temp_docx_file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + ' '
    return len(text.split())

# Function to convert Word document to CSV
def convert_doc_to_csv(input_docx, output_csv):
    doc = Document(input_docx)
    data = {'Text': [paragraph.text for paragraph in doc.paragraphs]}
    df = pd.DataFrame(data)
    df.to_csv(output_csv, index=False)

# Function to count words in CSV
def count_words_in_csv(csv_file):
    try:
        df = pd.read_csv(csv_file)
        total_words = 0
        for text in df['Text']:
            if isinstance(text, str):  # Check if text is a string
                total_words += len(text.split())
        return total_words
    except Exception as e:
        print("Error:", e)
        return 0  # Return 0 words if there's an error

def word_counter(comment):
    # extract text from the comment
    words = comment.split()
    word_count = len(words)
    return word_count

